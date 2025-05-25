from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout
from django.db.models import Count
from django.contrib.auth import get_user_model
from django.utils import timezone
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_POST
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from datetime import timedelta
from django.contrib import messages
from django.core.mail import send_mail
from django.conf import settings as django_settings
from functools import wraps
import os
import io
import docx
import PyPDF2
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile

from quiz.models import Quiz, Question, QuizAttempt
from curriculum.models import Curriculum, ClassLevel, Subject, Topic, SubTopic, Note, Attachment
from subscription.models import Subscription, SubscriptionPlan
from core.models import SystemConfiguration

User = get_user_model()


def admin_required(view_func):
    """Custom decorator to require admin authentication for admin views."""
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('my_admin:login')
        if not (request.user.is_staff or request.user.is_superuser):
            messages.error(request, "You don't have permission to access the admin area.")
            return redirect('core:dashboard')
        return view_func(request, *args, **kwargs)
    return wrapper


def admin_login(request):
    """Custom admin login view."""
    if request.user.is_authenticated and (request.user.is_staff or request.user.is_superuser):
        return redirect('my_admin:dashboard')

    if request.method == 'POST':
        email = request.POST.get('email')
        password = request.POST.get('password')

        if email and password:
            user = authenticate(request, username=email, password=password)
            if user is not None:
                if user.is_staff or user.is_superuser:
                    login(request, user)
                    next_url = request.GET.get('next', 'my_admin:dashboard')
                    return redirect(next_url)
                else:
                    messages.error(request, "You don't have permission to access the admin area.")
            else:
                messages.error(request, "Invalid email or password.")
        else:
            messages.error(request, "Please provide both email and password.")

    return render(request, 'my_admin/login.html')


def admin_logout(request):
    """Custom admin logout view."""
    logout(request)
    messages.success(request, "You have been logged out successfully.")
    return redirect('my_admin:login')


@admin_required
def help_center(request):
    """Admin help center view."""
    return render(request, 'my_admin/help.html')


@admin_required
def dashboard(request):
    """Main admin dashboard view."""
    # Get counts for various models
    total_users = User.objects.count()
    total_quizzes = Quiz.objects.count()
    total_questions = Question.objects.count()
    total_quiz_attempts = QuizAttempt.objects.count()

    # Get recent users
    recent_users = User.objects.order_by('-date_joined')[:5]

    # Get recent quiz attempts
    recent_quiz_attempts = QuizAttempt.objects.select_related('user', 'quiz').order_by('-started_at')[:5]

    # Get subscription stats
    subscription_stats = Subscription.objects.values('plan__plan_type').annotate(count=Count('id'))

    # Get user registration stats for the last 7 days
    today = timezone.now().date()
    seven_days_ago = today - timedelta(days=7)
    user_registration_stats = []

    for i in range(7):
        day = seven_days_ago + timedelta(days=i)
        count = User.objects.filter(date_joined__date=day).count()
        user_registration_stats.append({
            'day': day.strftime('%a'),
            'count': count
        })

    context = {
        'total_users': total_users,
        'total_quizzes': total_quizzes,
        'total_questions': total_questions,
        'total_quiz_attempts': total_quiz_attempts,
        'recent_users': recent_users,
        'recent_quiz_attempts': recent_quiz_attempts,
        'subscription_stats': subscription_stats,
        'user_registration_stats': user_registration_stats,
    }

    return render(request, 'my_admin/dashboard.html', context)

@admin_required
def quiz_management(request):
    """Quiz management view with search and filter capabilities."""
    # Get filter parameters from request
    curriculum_id = request.GET.get('curriculum')
    class_level_id = request.GET.get('class_level')
    subject_id = request.GET.get('subject')
    topic_id = request.GET.get('topic')
    status = request.GET.get('status')
    search_query = request.GET.get('search', '').strip()

    # Optimized query for speed and memory balance
    quizzes = Quiz.objects.select_related(
        'curriculum', 'class_level', 'subject', 'topic'
    ).order_by('-created_at')

    # Apply filters if provided
    if curriculum_id:
        quizzes = quizzes.filter(curriculum_id=curriculum_id)

    if class_level_id:
        quizzes = quizzes.filter(class_level_id=class_level_id)

    if subject_id:
        quizzes = quizzes.filter(subject_id=subject_id)

    if topic_id:
        quizzes = quizzes.filter(topic_id=topic_id)

    if status:
        is_active = status == 'active'
        quizzes = quizzes.filter(is_active=is_active)

    # Apply search if provided
    if search_query:
        quizzes = quizzes.filter(title__icontains=search_query)

    # PAGINATION: 20 quizzes per page
    paginator = Paginator(quizzes, 20)
    page = request.GET.get('page', 1)

    try:
        quizzes = paginator.page(page)
    except PageNotAnInteger:
        quizzes = paginator.page(1)
    except EmptyPage:
        quizzes = paginator.page(paginator.num_pages)

    # Get filter options for dropdowns - balanced for speed and memory
    curricula = Curriculum.objects.filter(is_active=True).order_by('name')
    class_levels = ClassLevel.objects.filter(is_active=True).select_related('curriculum').order_by('curriculum__name', 'level_order')
    subjects = Subject.objects.filter(is_active=True).order_by('name')
    topics = Topic.objects.filter(is_active=True).select_related('subject').order_by('subject__name', 'name')

    context = {
        'quizzes': quizzes,
        'curricula': curricula,
        'class_levels': class_levels,
        'subjects': subjects,
        'topics': topics,
        'selected_curriculum': curriculum_id,
        'selected_class_level': class_level_id,
        'selected_subject': subject_id,
        'selected_topic': topic_id,
        'selected_status': status,
        'search_query': search_query,
    }

    return render(request, 'my_admin/quiz_management.html', context)

@admin_required
def user_management(request):
    """Ultra-lightweight user management for 512MB memory limit."""
    try:
        # EXTREME memory optimization for Render's 512MB limit
        users_list = User.objects.only('id', 'email', 'first_name', 'last_name', 'is_active', 'is_staff', 'date_joined', 'last_login').order_by('-date_joined')

        # Pagination with smaller page size for memory
        paginator = Paginator(users_list, 10)  # Reduced to 10 users per page
        page = request.GET.get('page', 1)

        try:
            users = paginator.page(page)
        except (PageNotAnInteger, EmptyPage):
            users = paginator.page(1)

        # Simple stats without complex queries
        context = {
            'users': users,
            'total_users': 0,  # Disable to save memory
            'active_users_count': 0,  # Disable to save memory
            'premium_users_count': 0,  # Disable to save memory
            'new_users_today': 0,  # Disable to save memory
        }

        return render(request, 'my_admin/user_management.html', context)

    except Exception as e:
        # Fallback for any memory issues
        context = {
            'users': [],
            'total_users': 0,
            'active_users_count': 0,
            'premium_users_count': 0,
            'new_users_today': 0,
            'error': 'Memory optimization in progress'
        }
        return render(request, 'my_admin/user_management.html', context)


@admin_required
def user_details_api(request, user_id):
    """API endpoint to get user details."""
    try:
        user = get_object_or_404(User, id=user_id)

        # Get user's subscription info
        subscription_info = None
        if user.subscriptions.exists():
            latest_subscription = user.subscriptions.first()
            subscription_info = latest_subscription.plan.get_plan_type_display()

        data = {
            'id': user.id,
            'email': user.email,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'full_name': user.get_full_name(),
            'phone_number': getattr(user, 'phone_number', ''),
            'is_active': user.is_active,
            'is_staff': user.is_staff,
            'date_joined': user.date_joined.isoformat(),
            'last_login': user.last_login.isoformat() if user.last_login else None,
            'subscription': subscription_info,
        }

        return JsonResponse(data)
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@admin_required
def user_update_api(request, user_id):
    """API endpoint to update user details."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    try:
        import json
        user = get_object_or_404(User, id=user_id)
        data = json.loads(request.body)

        # Update user fields
        user.email = data.get('email', user.email)
        user.first_name = data.get('first_name', user.first_name)
        user.last_name = data.get('last_name', user.last_name)
        user.is_active = data.get('is_active', user.is_active)
        user.is_staff = data.get('is_staff', user.is_staff)

        # Update phone number if the field exists
        if hasattr(user, 'phone_number'):
            user.phone_number = data.get('phone_number', user.phone_number)

        user.save()

        return JsonResponse({'success': True, 'message': 'User updated successfully'})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@admin_required
def user_toggle_status_api(request, user_id):
    """API endpoint to toggle user active status."""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    try:
        import json
        user = get_object_or_404(User, id=user_id)
        data = json.loads(request.body)

        user.is_active = data.get('activate', not user.is_active)
        user.save()

        status = 'activated' if user.is_active else 'deactivated'
        return JsonResponse({'success': True, 'message': f'User {status} successfully'})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@admin_required
def user_delete_api(request, user_id):
    """API endpoint to delete a user."""
    if request.method != 'DELETE':
        return JsonResponse({'error': 'Method not allowed'}, status=405)

    try:
        user = get_object_or_404(User, id=user_id)

        # Don't allow deleting superusers or the current user
        if user.is_superuser:
            return JsonResponse({'error': 'Cannot delete superuser'}, status=400)

        if user.id == request.user.id:
            return JsonResponse({'error': 'Cannot delete yourself'}, status=400)

        user_email = user.email
        user.delete()

        return JsonResponse({'success': True, 'message': f'User {user_email} deleted successfully'})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@admin_required
def user_export(request):
    """Export users to CSV."""
    import csv
    from django.http import HttpResponse

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="users_export.csv"'

    writer = csv.writer(response)
    writer.writerow(['Email', 'First Name', 'Last Name', 'Phone', 'Active', 'Staff', 'Date Joined', 'Last Login', 'Subscription'])

    users = User.objects.prefetch_related('subscriptions').order_by('-date_joined')

    for user in users:
        subscription = 'Free'
        if user.subscriptions.exists():
            subscription = user.subscriptions.first().plan.get_plan_type_display()

        writer.writerow([
            user.email,
            user.first_name,
            user.last_name,
            getattr(user, 'phone_number', ''),
            'Yes' if user.is_active else 'No',
            'Yes' if user.is_staff else 'No',
            user.date_joined.strftime('%Y-%m-%d'),
            user.last_login.strftime('%Y-%m-%d') if user.last_login else 'Never',
            subscription
        ])

    return response

@admin_required
def curriculum_management(request):
    """Curriculum management view with memory optimization."""
    # MEMORY OPTIMIZATION: Limit results to prevent memory overload
    curricula = Curriculum.objects.all()[:20]
    class_levels = ClassLevel.objects.select_related('curriculum').all()[:50]
    subjects = Subject.objects.select_related('curriculum', 'class_level').all()[:100]
    topics = Topic.objects.select_related('subject', 'subject__curriculum', 'subject__class_level').all()[:200]

    context = {
        'curricula': curricula,
        'class_levels': class_levels,
        'subjects': subjects,
        'topics': topics,
    }

    return render(request, 'my_admin/curriculum_management.html', context)

@admin_required
def subscription_management(request):
    """Subscription management view."""
    subscriptions = Subscription.objects.all().select_related('user', 'plan').order_by('-created_at')

    # Get all users and plans for the add subscription form
    users = User.objects.all().order_by('email')
    all_plans = SubscriptionPlan.objects.all().order_by('plan_type', 'price')

    # Get subscription stats by plan type
    plan_type_stats = {}
    for plan_type, plan_name in SubscriptionPlan.PLAN_TYPE_CHOICES:
        count = subscriptions.filter(plan__plan_type=plan_type).count()
        plan_type_stats[plan_name] = count

    # Get subscription stats by status
    status_stats = {}
    for status_code, status_name in Subscription.STATUS_CHOICES:
        count = subscriptions.filter(status=status_code).count()
        status_stats[status_name] = count

    # Add count for active subscriptions based on is_active property
    active_count = sum(1 for sub in subscriptions if sub.is_active)
    inactive_count = len(subscriptions) - active_count

    # Prepare data for charts
    plan_labels = list(plan_type_stats.keys())
    plan_counts = list(plan_type_stats.values())

    status_labels = list(status_stats.keys())
    status_counts = list(status_stats.values())

    # Add active/inactive stats
    active_stats = {
        'Active': active_count,
        'Inactive': inactive_count
    }
    active_labels = list(active_stats.keys())
    active_counts = list(active_stats.values())

    context = {
        'subscriptions': subscriptions,
        'users': users,
        'all_plans': all_plans,
        'status_choices': Subscription.STATUS_CHOICES,
        'plan_type_stats': plan_type_stats,
        'status_stats': status_stats,
        'active_stats': active_stats,
        'plan_labels': plan_labels,
        'plan_counts': plan_counts,
        'status_labels': status_labels,
        'status_counts': status_counts,
        'active_labels': active_labels,
        'active_counts': active_counts,
    }

    return render(request, 'my_admin/subscription_management.html', context)


@admin_required
def add_subscription(request):
    """Add a new subscription."""
    if request.method == 'POST':
        user_id = request.POST.get('user')
        plan_id = request.POST.get('plan')
        status = request.POST.get('status')
        duration = int(request.POST.get('duration', 365))
        auto_renew = 'auto_renew' in request.POST

        if user_id and plan_id:
            user = get_object_or_404(User, id=user_id)
            plan = get_object_or_404(SubscriptionPlan, id=plan_id)

            # Calculate start and end dates
            start_date = timezone.now()
            end_date = start_date + timedelta(days=duration)

            # Create the subscription
            subscription = Subscription.objects.create(
                user=user,
                plan=plan,
                status=status,
                start_date=start_date,
                end_date=end_date,
                auto_renew=auto_renew
            )

            messages.success(request, f"Subscription for {user.email} has been created successfully.")
        else:
            messages.error(request, "Failed to create subscription. Please check the form and try again.")

    return redirect('my_admin:subscription_management')


@admin_required
def edit_subscription(request, subscription_id):
    """Edit a subscription."""
    subscription = get_object_or_404(Subscription, id=subscription_id)

    if request.method == 'POST':
        # Get form data
        plan_id = request.POST.get('plan')
        status = request.POST.get('status')
        auto_renew = request.POST.get('auto_renew') == 'on'

        # Update subscription
        if plan_id:
            subscription.plan_id = plan_id

        subscription.status = status
        subscription.auto_renew = auto_renew
        subscription.save()

        messages.success(request, f"Subscription for {subscription.user.email} has been updated.")
        return redirect('my_admin:subscription_management')

    # Get all available plans
    plans = subscription.plan.__class__.objects.all()

    context = {
        'subscription': subscription,
        'plans': plans,
        'status_choices': Subscription.STATUS_CHOICES,
    }

    return render(request, 'my_admin/edit_subscription.html', context)


@admin_required
def renew_subscription(request, subscription_id):
    """Renew a subscription."""
    subscription = get_object_or_404(Subscription, id=subscription_id)

    if request.method == 'POST':
        # Get form data
        days = int(request.POST.get('days', 30))

        # Update subscription end date
        subscription.end_date = subscription.end_date + timedelta(days=days)

        # If subscription is expired or canceled, reactivate it
        if subscription.status in ['expired', 'canceled']:
            subscription.status = 'active'
            subscription.canceled_at = None

        subscription.save()

        messages.success(request, f"Subscription for {subscription.user.email} has been renewed for {days} days.")
        return redirect('my_admin:subscription_management')

    context = {
        'subscription': subscription,
    }

    return render(request, 'my_admin/renew_subscription.html', context)


# Curriculum Management Views
@admin_required
@require_POST
def add_curriculum(request):
    """Add a new curriculum."""
    name = request.POST.get('name')
    code = request.POST.get('code')
    country = request.POST.get('country')
    description = request.POST.get('description', '')
    is_active = request.POST.get('is_active') == 'on'

    if name and code and country:
        Curriculum.objects.create(
            name=name,
            code=code,
            country=country,
            description=description,
            is_active=is_active
        )

    return redirect('my_admin:curriculum_management')


@admin_required
@require_POST
def add_class_level(request):
    """Add a new class level."""
    name = request.POST.get('name')
    curriculum_id = request.POST.get('curriculum')
    order = request.POST.get('order', 1)
    is_active = request.POST.get('is_active') == 'on'

    if name and curriculum_id:
        curriculum = get_object_or_404(Curriculum, id=curriculum_id)
        ClassLevel.objects.create(
            name=name,
            curriculum=curriculum,
            order=order,
            is_active=is_active
        )

    return redirect('my_admin:curriculum_management')


@admin_required
@require_POST
def add_subject(request):
    """Add a new subject."""
    name = request.POST.get('name')
    curriculum_id = request.POST.get('curriculum')
    class_level_id = request.POST.get('class_level')
    description = request.POST.get('description', '')
    is_active = request.POST.get('is_active') == 'on'

    if name and curriculum_id and class_level_id:
        curriculum = get_object_or_404(Curriculum, id=curriculum_id)
        class_level = get_object_or_404(ClassLevel, id=class_level_id)
        Subject.objects.create(
            name=name,
            curriculum=curriculum,
            class_level=class_level,
            description=description,
            is_active=is_active
        )

    return redirect('my_admin:curriculum_management')


@admin_required
@require_POST
def add_topic(request):
    """Add a new topic."""
    name = request.POST.get('name')
    subject_id = request.POST.get('subject')
    order = request.POST.get('order', 1)
    description = request.POST.get('description', '')
    is_active = request.POST.get('is_active') == 'on'

    if name and subject_id:
        subject = get_object_or_404(Subject, id=subject_id)
        Topic.objects.create(
            name=name,
            subject=subject,
            order=order,
            description=description,
            is_active=is_active
        )

    return redirect('my_admin:curriculum_management')


@admin_required
@require_POST
def edit_curriculum(request, curriculum_id):
    """Edit an existing curriculum."""
    curriculum = get_object_or_404(Curriculum, id=curriculum_id)

    name = request.POST.get('name')
    code = request.POST.get('code')
    country = request.POST.get('country')
    description = request.POST.get('description', '')
    is_active = request.POST.get('is_active') == 'on'

    if name and code and country:
        curriculum.name = name
        curriculum.code = code
        curriculum.country = country
        curriculum.description = description
        curriculum.is_active = is_active
        curriculum.save()

    return redirect('my_admin:curriculum_management')


# Site Settings Views
@admin_required
def site_settings_general(request):
    """View for general site settings."""
    settings = SystemConfiguration.get_settings()

    if request.method == 'POST':
        # Update settings
        settings.site_name = request.POST.get('site_name', settings.site_name)
        settings.site_description = request.POST.get('site_description', settings.site_description)
        settings.maintenance_mode = 'maintenance_mode' in request.POST
        settings.enforce_single_session = 'enforce_single_session' in request.POST

        try:
            settings.session_timeout_minutes = int(request.POST.get('session_timeout_minutes', settings.session_timeout_minutes))
        except (ValueError, TypeError):
            pass

        settings.updated_by = request.user
        settings.save()

        messages.success(request, "General settings updated successfully.")
        return redirect('my_admin:site_settings_general')

    return render(request, 'admin/site_settings/general.html', {
        'settings': settings,
        'active_tab': 'general',
    })


@admin_required
def site_settings_email(request):
    """View for email settings."""
    settings = SystemConfiguration.get_settings()

    if request.method == 'POST':
        # Update settings
        settings.smtp_host = request.POST.get('smtp_host', settings.smtp_host)
        settings.smtp_username = request.POST.get('smtp_username', settings.smtp_username)
        settings.smtp_password = request.POST.get('smtp_password', settings.smtp_password)
        settings.default_from_email = request.POST.get('default_from_email', settings.default_from_email)
        settings.smtp_use_tls = 'smtp_use_tls' in request.POST

        try:
            settings.smtp_port = int(request.POST.get('smtp_port', settings.smtp_port))
        except (ValueError, TypeError):
            pass

        settings.updated_by = request.user
        settings.save()

        messages.success(request, "Email settings updated successfully.")
        return redirect('my_admin:site_settings_email')

    return render(request, 'admin/site_settings/email.html', {
        'settings': settings,
        'active_tab': 'email',
    })


@admin_required
def site_settings_test_email(request):
    """Send a test email."""
    if request.method == 'POST':
        recipient = request.POST.get('test_email')
        subject = request.POST.get('test_subject', 'Test Email from EduMore360')
        message = request.POST.get('test_message', 'This is a test email from EduMore360.')

        try:
            send_mail(
                subject,
                message,
                django_settings.DEFAULT_FROM_EMAIL,
                [recipient],
                fail_silently=False,
            )
            messages.success(request, f"Test email sent to {recipient} successfully.")
        except Exception as e:
            messages.error(request, f"Failed to send test email: {str(e)}")

    return redirect('my_admin:site_settings_email')


@admin_required
def site_settings_quiz(request):
    """View for quiz settings."""
    settings = SystemConfiguration.get_settings()

    if request.method == 'POST':
        # Update settings
        settings.quiz_show_immediate_feedback = 'quiz_show_immediate_feedback' in request.POST
        settings.quiz_randomize_questions = 'quiz_randomize_questions' in request.POST
        settings.quiz_randomize_choices = 'quiz_randomize_choices' in request.POST

        try:
            settings.quiz_feedback_time = int(request.POST.get('quiz_feedback_time', settings.quiz_feedback_time))
        except (ValueError, TypeError):
            pass

        settings.updated_by = request.user
        settings.save()

        messages.success(request, "Quiz settings updated successfully.")
        return redirect('my_admin:site_settings_quiz')

    return render(request, 'admin/site_settings/quiz.html', {
        'settings': settings,
        'active_tab': 'quiz',
    })


@admin_required
def site_settings_payment(request):
    """View for payment settings."""
    settings = SystemConfiguration.get_settings()

    if request.method == 'POST':
        # Update settings
        settings.paystack_secret_key = request.POST.get('paystack_secret_key', settings.paystack_secret_key)
        settings.paystack_public_key = request.POST.get('paystack_public_key', settings.paystack_public_key)
        settings.currency = request.POST.get('currency', settings.currency)

        settings.updated_by = request.user
        settings.save()

        messages.success(request, "Payment settings updated successfully.")
        return redirect('my_admin:site_settings_payment')

    return render(request, 'admin/site_settings/payment.html', {
        'settings': settings,
        'active_tab': 'payment',
    })


# Notes Management Views
@admin_required
def notes_management(request):
    """View for managing notes with memory optimization."""
    # MEMORY OPTIMIZATION: Limit notes to prevent memory overload
    notes = Note.objects.all().select_related('topic', 'subtopic', 'created_by').order_by('-updated_at')[:100]
    curricula = Curriculum.objects.filter(is_active=True)[:20]

    context = {
        'notes': notes,
        'curricula': curricula,
        'file_type_choices': Note.FILE_TYPE_CHOICES,
    }

    return render(request, 'my_admin/notes_management.html', context)


@admin_required
def add_note_page(request):
    """View for adding a new note."""
    curricula = Curriculum.objects.filter(is_active=True)

    context = {
        'curricula': curricula,
        'file_type_choices': Note.FILE_TYPE_CHOICES,
    }

    return render(request, 'my_admin/add_note.html', context)


@admin_required
def add_note(request):
    """Add a new note."""
    if request.method != 'POST':
        return redirect('my_admin:add_note_page')

    title = request.POST.get('title')
    topic_id = request.POST.get('topic')
    subtopic_id = request.POST.get('subtopic') if request.POST.get('subtopic') else None
    content = request.POST.get('content', '')
    summary = request.POST.get('summary', '')
    is_premium = request.POST.get('is_premium') == 'on'
    is_published = request.POST.get('is_published') == 'on'
    allow_download = request.POST.get('allow_download') == 'on'
    order = request.POST.get('order', 0)
    file_type = request.POST.get('file_type', 'none')

    if not title or not topic_id:
        messages.error(request, "Failed to create note. Title and Topic are required.")
        return redirect('my_admin:add_note_page')

    topic = get_object_or_404(Topic, id=topic_id)
    subtopic = get_object_or_404(SubTopic, id=subtopic_id) if subtopic_id else None

    # Create the note
    # Clean the content to ensure it's properly formatted HTML
    cleaned_content = content.strip() if content else ""

    # Debug content
    if not cleaned_content:
        messages.warning(request, "Note content appears to be empty. Please check the editor.")

    # Create the note
    note = Note.objects.create(
        title=title,
        topic=topic,
        subtopic=subtopic,
        content=cleaned_content,
        summary=summary,
        is_premium=is_premium,
        is_published=is_published,
        allow_download=allow_download,
        order=order,
        file_type=file_type,
        created_by=request.user,
        updated_by=request.user
    )

    # Handle file upload if provided
    uploaded_file = request.FILES.get('document')
    if uploaded_file and file_type != 'none':
        # Process the uploaded file based on type
        if file_type == 'pdf' and uploaded_file.name.lower().endswith('.pdf'):
            note.pdf_document = uploaded_file

            # Try to extract text from PDF
            try:
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                extracted_text = ""
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    extracted_text += page.extract_text() + "\n\n"
                note.extracted_text = extracted_text
            except Exception as e:
                messages.warning(request, f"Could not extract text from PDF: {str(e)}")

        elif file_type == 'doc' and uploaded_file.name.lower().endswith(('.doc', '.docx')):
            note.doc_document = uploaded_file

            # Try to extract text from Word document
            try:
                # Save the file temporarily
                temp_path = default_storage.save('temp_docs/' + uploaded_file.name, ContentFile(uploaded_file.read()))
                uploaded_file.seek(0)  # Reset file pointer

                # Extract text
                doc = docx.Document(default_storage.path(temp_path))
                extracted_text = ""
                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"
                note.extracted_text = extracted_text

                # Clean up
                default_storage.delete(temp_path)
            except Exception as e:
                messages.warning(request, f"Could not extract text from Word document: {str(e)}")

        elif file_type == 'ppt' and uploaded_file.name.lower().endswith(('.ppt', '.pptx')):
            note.ppt_document = uploaded_file
            messages.info(request, "Text extraction from PowerPoint files is not supported.")

        else:
            note.other_document = uploaded_file
            note.file_type = 'other'
            messages.info(request, "Text extraction is not supported for this file type.")

        note.save()

    messages.success(request, f"Note '{note.title}' created successfully.")

    # Redirect to edit page for further editing if needed
    return redirect('my_admin:edit_note', note_id=note.id)


@admin_required
def edit_note(request, note_id):
    """Edit an existing note."""
    note = get_object_or_404(Note, id=note_id)

    if request.method == 'POST':
        title = request.POST.get('title')
        topic_id = request.POST.get('topic')
        subtopic_id = request.POST.get('subtopic') if request.POST.get('subtopic') else None
        content = request.POST.get('content', '')
        summary = request.POST.get('summary', '')
        is_premium = request.POST.get('is_premium') == 'on'
        is_published = request.POST.get('is_published') == 'on'
        order = request.POST.get('order', 0)
        allow_download = request.POST.get('allow_download') == 'on'

        if title and topic_id:
            topic = get_object_or_404(Topic, id=topic_id)
            subtopic = get_object_or_404(SubTopic, id=subtopic_id) if subtopic_id else None

            # Clean the content to ensure it's properly formatted HTML
            cleaned_content = content.strip() if content else ""

            # Debug content
            if not cleaned_content:
                messages.warning(request, "Note content appears to be empty. Please check the editor.")

            note.title = title
            note.topic = topic
            note.subtopic = subtopic
            note.content = cleaned_content
            note.summary = summary
            note.is_premium = is_premium
            note.is_published = is_published
            note.order = order
            note.allow_download = allow_download
            note.updated_by = request.user
            note.save()

            messages.success(request, f"Note '{note.title}' updated successfully.")
            return redirect('my_admin:notes_management')
        else:
            messages.error(request, "Failed to update note. Please check the form and try again.")

    # Get all curricula, class levels, subjects, topics, and subtopics for the form
    curricula = Curriculum.objects.filter(is_active=True)
    class_levels = ClassLevel.objects.filter(curriculum=note.topic.subject.curriculum, is_active=True)
    subjects = Subject.objects.filter(class_level=note.topic.subject.class_level, is_active=True)
    topics = Topic.objects.filter(subject=note.topic.subject, is_active=True)
    subtopics = SubTopic.objects.filter(topic=note.topic, is_active=True) if note.topic else []

    context = {
        'note': note,
        'curricula': curricula,
        'class_levels': class_levels,
        'subjects': subjects,
        'topics': topics,
        'subtopics': subtopics,
        'file_type_choices': Note.FILE_TYPE_CHOICES,
    }

    return render(request, 'my_admin/edit_note.html', context)


@admin_required
@require_POST
def delete_note(request, note_id):
    """Delete a note."""
    note = get_object_or_404(Note, id=note_id)
    title = note.title
    note.delete()

    messages.success(request, f"Note '{title}' deleted successfully.")
    return redirect('my_admin:notes_management')


@admin_required
@require_POST
def upload_note_file(request, note_id):
    """Upload a file for a note."""
    note = get_object_or_404(Note, id=note_id)
    file_type = request.POST.get('file_type')
    uploaded_file = request.FILES.get('document')

    if not uploaded_file or not file_type:
        messages.error(request, "No file uploaded or file type not specified.")
        return redirect('my_admin:edit_note', note_id=note.id)

    # Clear any existing files
    if note.pdf_document:
        note.pdf_document.delete()
    if note.ppt_document:
        note.ppt_document.delete()
    if note.doc_document:
        note.doc_document.delete()
    if note.other_document:
        note.other_document.delete()

    # Save the new file based on type
    note.file_type = file_type

    try:
        # Extract text from document if possible
        extracted_text = ""

        if file_type == 'pdf' and uploaded_file.name.lower().endswith('.pdf'):
            note.pdf_document = uploaded_file

            # Try to extract text from PDF
            try:
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    extracted_text += page.extract_text() + "\n\n"
            except Exception as e:
                messages.warning(request, f"Could not extract text from PDF: {str(e)}")

        elif file_type == 'doc' and uploaded_file.name.lower().endswith(('.doc', '.docx')):
            note.doc_document = uploaded_file

            # Try to extract text from Word document
            try:
                # Save the file temporarily
                temp_path = default_storage.save('temp_docs/' + uploaded_file.name, ContentFile(uploaded_file.read()))
                uploaded_file.seek(0)  # Reset file pointer

                # Extract text
                doc = docx.Document(default_storage.path(temp_path))
                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"

                # Clean up
                default_storage.delete(temp_path)
            except Exception as e:
                messages.warning(request, f"Could not extract text from Word document: {str(e)}")

        elif file_type == 'ppt' and uploaded_file.name.lower().endswith(('.ppt', '.pptx')):
            note.ppt_document = uploaded_file
            messages.info(request, "Text extraction from PowerPoint files is not supported.")

        else:
            note.other_document = uploaded_file
            note.file_type = 'other'
            messages.info(request, "Text extraction is not supported for this file type.")

        # Save extracted text if any
        if extracted_text:
            note.extracted_text = extracted_text
            messages.success(request, "Text successfully extracted from document.")

        note.updated_by = request.user
        note.save()

        messages.success(request, f"File uploaded successfully for note '{note.title}'.")
    except Exception as e:
        messages.error(request, f"Error uploading file: {str(e)}")

    return redirect('my_admin:edit_note', note_id=note.id)


@admin_required
def get_topics_for_notes(request):
    """AJAX endpoint to get topics for a subject."""
    subject_id = request.GET.get('subject')

    if subject_id:
        topics = Topic.objects.filter(subject_id=subject_id, is_active=True).values('id', 'name')
        return JsonResponse(list(topics), safe=False)

    return JsonResponse([], safe=False)


@admin_required
def get_subtopics_for_notes(request):
    """AJAX endpoint to get subtopics for a topic."""
    topic_id = request.GET.get('topic')

    if topic_id:
        subtopics = SubTopic.objects.filter(topic_id=topic_id, is_active=True).values('id', 'name')
        return JsonResponse(list(subtopics), safe=False)

    return JsonResponse([], safe=False)